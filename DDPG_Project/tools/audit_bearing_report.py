from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


REQUIRED_TERMS = [
    "16 kHz",
    "4 ms",
    "64",
    "50 ms",
    "800",
    "200 ms",
    "3200",
    "List_seq",
    "PER-DDPG",
    "MQTT",
    "边缘节点",
    "云端",
]

ALLOWED_FOG_CONTEXTS = [
    "论文原始网络模型包含云、雾、边三层",
]


def collect_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("path", type=Path)
    args = parser.parse_args()

    document = Document(args.path)
    text = collect_text(document)
    missing = [term for term in REQUIRED_TERMS if term not in text]
    headings = [p.text for p in document.paragraphs if p.style.name == "Heading 1"]
    fog_lines = [line for line in text.splitlines() if "雾" in line]
    unexpected_fog = [
        line for line in fog_lines if not any(context in line for context in ALLOWED_FOG_CONTEXTS)
    ]
    empty_tables = [index for index, table in enumerate(document.tables) if not table.rows]
    header_text = " ".join(p.text for p in document.sections[0].header.paragraphs).strip()
    footer_xml = document.sections[0].footer._element.xml

    problems = []
    if missing:
        problems.append(f"missing required terms: {missing}")
    if len(headings) < 8:
        problems.append(f"expected at least 8 Heading 1 sections, found {len(headings)}")
    if not document.tables:
        problems.append("expected at least one table")
    if empty_tables:
        problems.append(f"empty tables: {empty_tables}")
    if not header_text:
        problems.append("missing running header")
    if "PAGE" not in footer_xml:
        problems.append("missing PAGE field in footer")
    if unexpected_fog:
        problems.append(f"unexpected fog-layer wording: {unexpected_fog}")

    print(f"paragraphs={len(document.paragraphs)}")
    print(f"heading1={len(headings)}")
    print(f"tables={len(document.tables)}")
    print(f"fog_occurrences={len(fog_lines)}")
    for line in fog_lines:
        print(f"allowed_fog_context={line}")

    if problems:
        for problem in problems:
            print(f"ERROR: {problem}")
        raise SystemExit(1)
    print("AUDIT PASS")


if __name__ == "__main__":
    main()
