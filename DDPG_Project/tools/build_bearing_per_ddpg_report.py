from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


INK = "202A35"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5D6670"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
WHITE = "FFFFFF"
GOLD = "7A5A00"
RED = "9B1C1C"
BODY_FONT = "Calibri"
EAST_ASIA_FONT = "Microsoft YaHei"
CONTENT_WIDTH_DXA = 9360


def set_run_font(run, size=11, bold=False, color=INK, italic=False, east_asia=EAST_ASIA_FONT):
    run.font.name = BODY_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), BODY_FONT)
    rfonts.set(qn("w:hAnsi"), BODY_FONT)
    rfonts.set(qn("w:eastAsia"), east_asia)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        prevent_row_split(row)
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, value, end])
    suffix = paragraph.add_run(" 页")
    set_run_font(suffix, size=9, color=MUTED)


def configure_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    normal.paragraph_format.widow_control = True

    settings = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in settings.items():
        style = styles[name]
        style.font.name = BODY_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = BODY_FONT
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_running_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("工业轴承检测场景 | PER-DDPG 调度技术方案")
    set_run_font(run, size=9, bold=True, color=MUTED)
    footer = section.footer
    add_page_number(footer.paragraphs[0])


def add_paragraph(document, text, bold_prefix=None, italic=False, color=INK, keep=False):
    p = document.add_paragraph()
    p.paragraph_format.keep_together = keep
    if bold_prefix and text.startswith(bold_prefix):
        lead = p.add_run(bold_prefix)
        set_run_font(lead, bold=True, color=color)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest, color=color, italic=italic)
    else:
        run = p.add_run(text)
        set_run_font(run, color=color, italic=italic)
    return p


def add_bullet(document, text):
    p = document.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_number(document, text):
    p = document.add_paragraph(style="List Number")
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_callout(document, label, text, fill=PALE_BLUE, color=DARK_BLUE):
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.keep_together = True
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), color)
    p_bdr.append(left)
    p_pr.append(p_bdr)
    lead = p.add_run(f"{label}：")
    set_run_font(lead, bold=True, color=color)
    body = p.add_run(text)
    set_run_font(body, color=INK)


def add_table(document, headers, rows, widths_dxa, font_size=9.5):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, text in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, LIGHT)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=font_size, bold=True, color=DARK_BLUE)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, text in enumerate(values):
            if row_index % 2 == 1:
                set_cell_shading(cells[index], "FAFBFC")
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(text))
            set_run_font(r, size=font_size)
    set_table_geometry(table, widths_dxa)
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_title_block(document):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("技术分析与实施方案")
    set_run_font(run, size=10.5, bold=True, color=BLUE)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(5)
    title.paragraph_format.keep_with_next = True
    run = title.add_run("工业轴承检测场景下的 PER-DDPG\n任务调度与资源分配方案")
    set_run_font(run, size=23, bold=True, color=INK)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("基于两份会议材料与论文第 3 章算法思想的无中间计算层场景整合")
    set_run_font(run, size=12.5, color=MUTED)

    metadata = [
        ("分析对象", "《关于开会时提出的问题》《2026.7.5(2)》"),
        ("应用对象", "电机或旋转机械中的滚动轴承状态监测"),
        ("系统结构", "采集发送端—MQTT/调度器—边缘节点群—按需云端"),
        ("算法边界", "保留 DAG 排序 + PER-DDPG 卸载与连续资源分配思想"),
        ("文档日期", "2026 年 7 月 15 日"),
    ]
    for label, value in metadata:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        left = p.add_run(f"{label}：")
        set_run_font(left, bold=True, color=INK)
        right = p.add_run(value)
        set_run_font(right, color=INK)

    rule = document.add_paragraph()
    rule.paragraph_format.space_before = Pt(8)
    rule.paragraph_format.space_after = Pt(12)
    p_pr = rule._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_section_1(document):
    document.add_heading("1. 执行摘要", level=1)
    add_callout(
        document,
        "核心结论",
        "两份材料并不冲突：第一份补充了网络、MQTT、复核和弱网自治机制，第二份给出了轴承数据采集与接口参数。二者合并后，可将一次轴承状态分析建模为具有依赖关系的 DAG 任务，由算法 3.1 生成执行顺序，再由适配后的 PER-DDPG 为每个就绪子任务选择边缘节点或云端，并决定连续 CPU 资源比例。",
    )
    add_paragraph(document, "最终工业流程不是“采集后全部上云”，也不是“第一轮直接把所有任务一次性分好”。系统以 50 ms 微批量为基本输入，以 200 ms 滑动窗口形成诊断任务；调度器在每个任务步重新读取节点剩余资源和通信状态，执行当前策略，并把完成时间反馈写入经验池。经过多回合训练后，策略才逐步形成稳定的状态—动作映射。")
    add_paragraph(document, "当前方案明确去除论文原始网络模型中的中间计算层。工程侧只保留多个异构边缘节点和一个云端：常规数据优先在边缘完成，低置信度、指标矛盾、严重异常或复核冲突时才触发第二边缘节点或云端深度分析。这样既保留论文“拓扑排序 + 动态卸载 + 连续资源分配”的核心思想，又与轴承检测的实时性、弱网自治和安全需求一致。")


def add_section_2(document):
    document.add_heading("2. 两份文档的内容分析与联系", level=1)
    document.add_heading("2.1 《2026.7.5(2)》：定义数据入口与时间尺度", level=2)
    add_paragraph(document, "该文档解决的是“检测数据如何产生、如何打包、边缘节点拿到什么”的问题。场景被限定为滚动轴承状态监测，采集终端获取振动、电流、温度、转速和负载；发送器只封装原始或最新标量，不做故障判断；边缘节点承担缓存、预处理、特征提取和轻量模型推理。")
    add_paragraph(document, "文档已经给出较完整的时间链路：振动以 16 kHz 连续采样，每 4 ms 写入 64 个点；发送器每 50 ms 发送 800 点振动微批量，并附加最新电流、温度、转速和负载；边缘节点维护 200 ms、3200 点环形缓存，并每 50 ms 对最近窗口执行一次分析。这些参数构成后续 DAG 的任务到达周期、数据量和最早就绪时间。")
    add_paragraph(document, "它尚未解决的主要是：接口对新传感器和图像数据不够可扩展；没有定义多个边缘节点如何选择；没有明确弱网下发送降级、重复检测和上云条件；也没有给出资源分配和调度反馈闭环。")

    document.add_heading("2.2 《关于开会时提出的问题》：定义网络、路由和复核闭环", level=2)
    add_paragraph(document, "该文档将系统从单一接口扩展为可调度的工业物联网流程。它区分采集端到边缘的第一段通信和边缘到云端的第二段通信，指出网络状态不仅影响传输时延和数据完整性，还影响边缘是否上云、上传原始数据还是只上传 RMS、Peak 等关键特征。")
    add_paragraph(document, "文档在 HTTP、调度器直发和 MQTT 发布订阅三种方式中，推荐“MQTT Broker + 调度器控制路由”。发送器只负责采集和发布，调度器根据任务、节点和网络状态选择目标主题，边缘节点只处理分配给自己的任务。新增节点时无需修改发送器，职责边界更清楚。")
    add_paragraph(document, "它还提出两级判定：先选择一个合适的边缘节点初检；若出现低置信度、正常/异常边界、指标矛盾或严重异常，再调度另一个边缘节点复核；复核结果冲突时上传云端。该机制把“多个节点重复判断”从常态广播改成条件触发，降低通信和计算浪费。")

    document.add_heading("2.3 两份文档的互补关系", level=2)
    add_table(
        document,
        ["维度", "《2026.7.5(2)》", "《关于开会时提出的问题》", "合并后的作用"],
        [
            ("对象", "轴承多传感数据", "网络与多节点调度", "形成具体可运行的轴承检测系统"),
            ("时间尺度", "4 ms、50 ms、200 ms", "弱网、复核、上云时机", "把实时采集和动态决策连接起来"),
            ("通信", "给出微批量 JSON", "推荐 MQTT 路由", "定义统一消息信封和主题路由"),
            ("边缘职责", "缓存、预处理、特征、推理", "初检、复核、本地自治", "形成边缘闭环与分工"),
            ("未覆盖部分", "缺少调度算法", "缺少具体轴承任务参数", "由 DAG + PER-DDPG 补齐"),
        ],
        [1200, 2400, 2700, 3060],
        font_size=9,
    )
    add_callout(document, "统一后的任务定义", "一个 task 对应同一设备、同一 200 ms 分析窗口的诊断任务；task 内可包含振动、温度、电流、转速、负载等相关 packet。packet 可以并行预处理，但融合、诊断和告警必须遵守依赖关系。")


def add_section_3(document):
    document.add_heading("3. 与论文 PER-DDPG 算法的联系", level=1)
    add_paragraph(document, "论文第 3 章解决的是具有严格拓扑依赖的任务卸载问题，目标是在满足节点资源约束的情况下最小化总任务完成时间。其方法不是单个神经网络直接输出全部结果，而是分成两个阶段。")
    document.add_heading("3.1 阶段一：算法 3.1 生成 List_seq", level=2)
    add_paragraph(document, "首先用 DAG 表达子任务前驱关系。对每个子任务计算包含自身平均处理时间和后继关键路径开销的静态优先级 rank_i，再根据全局容忍时延、虚拟当前时间和剩余关键路径开销计算 urgency_i，最终按权重 W_i 综合二者。在每一步，只从全部前驱已完成的就绪任务中选择 W_i 最大者，生成满足拓扑约束的线性序列 List_seq。")
    add_callout(document, "重要边界", "List_seq 是启发式优先调度序列，不是全局最优证明。它解决“当前应该先处理哪个已就绪任务”，但不决定该任务具体放在哪个节点，也不决定分配多少 CPU。")

    document.add_heading("3.2 阶段二：PER-DDPG 选择节点和资源比例", level=2)
    add_paragraph(document, "智能体严格按 List_seq 逐项决策。状态 s_t 包含当前任务特征 I_t、候选计算节点剩余资源 CR_t 和通信条件 R_t；动作 a_t 由节点选择概率 A_t 与连续资源比例 F_t 构成；环境执行动作后计算传输、等待和执行时间，以当前总完成时刻的负值作为奖励。")
    add_paragraph(document, "Actor 输出动作，Critic 估计该动作的长期价值；目标网络通过软更新稳定训练。优先经验回放 PER 按 TD 误差提高关键经验的采样概率，自适应探索噪声在训练早期扩大探索、后期逐步减弱。论文通过这些机制提高复杂动作空间中的样本利用率和收敛稳定性。")

    document.add_heading("3.3 在轴承场景中的结构适配", level=2)
    add_paragraph(document, "论文原始网络模型包含云、雾、边三层。当前轴承方案明确取消中间计算层，因此动作候选不再包含该层节点，而改为“当前可用边缘节点集合 + 云端”。这是一项场景结构适配：保留状态—动作—奖励、Actor-Critic、PER 和探索机制，但重新定义候选节点索引及对应通信速率。")
    add_table(
        document,
        ["论文算法概念", "轴承场景映射", "说明"],
        [
            ("DAG 子任务", "接收校验、预处理、特征、初检、复核、云端分析、告警", "由数据依赖和业务条件确定前驱关系"),
            ("I_t", "数据量、计算密度、最小 CPU、最早就绪时间", "可由窗口与子任务类型生成"),
            ("CR_t", "各边缘节点与云端剩余 CPU", "反映排队和当前负载"),
            ("R_t", "采集端/当前数据位置到候选节点的传输速率", "反映弱网和链路差异"),
            ("A_t", "边缘节点 E1…Ek 或云端的选择概率", "取最大概率对应的可行节点"),
            ("F_t", "目标节点剩余 CPU 的请求比例", "执行时需满足最小 CPU 与容量约束"),
            ("奖励", "负的当前总完成时刻", "保持最小化任务完成时间的论文目标"),
            ("PER", "优先回放高 TD 误差的调度经验", "重点学习拥塞、弱网、冲突等高价值状态"),
        ],
        [1900, 3500, 3960],
        font_size=9,
    )

    document.add_heading("3.4 为什么第一轮决策不等于已经学会", level=2)
    add_paragraph(document, "第一轮确实可以产生一套完整分配方案，但“能够输出”与“已经学到较优策略”是两回事。未训练 Actor 的参数近似随机，第一轮动作只是随机初始化、探索噪声和当前状态共同作用的结果；它没有见过相同或相近状态下其他节点与资源比例的后果，也无法判断长期 makespan 是否更小。")
    add_bullet(document, "一次状态只告诉系统当前任务、资源和通信条件，不提供所有可能动作的真实结果。")
    add_bullet(document, "环境只执行一个动作，其他动作属于反事实，需要通过后续探索或价值估计比较。")
    add_bullet(document, "节点负载、通信速率、任务拓扑和到达时刻会变化，固定规则无法覆盖全部组合。")
    add_bullet(document, "Critic 需要大量转移样本学习长期回报；Actor 再依据 Critic 梯度修正节点概率和资源比例。")
    add_bullet(document, "训练后的策略仍不是数学意义上的全局最优保证，而是在训练分布下获得的近似优良策略。")


def add_section_4(document):
    document.add_heading("4. 无中间计算层的工业轴承检测架构", level=1)
    add_paragraph(document, "系统采用“采集发送端—MQTT Broker/调度器—边缘节点群—按需云端”架构。采集与诊断解耦，数据路由与故障判断解耦，常规任务与复核任务解耦。")
    add_table(
        document,
        ["组件", "核心职责", "明确不负责的事项"],
        [
            ("轴承与传感器", "采集振动、电流、温度、转速、负载", "不做故障判断和节点选择"),
            ("一体化采集发送器", "连续缓存、微批量封装、序号与时间戳、MQTT 发布", "不依据负载自行选择边缘节点"),
            ("MQTT Broker", "主题订阅、消息转发、QoS 与断线重连", "不执行调度策略和诊断模型"),
            ("调度器", "构造 DAG、生成 List_seq、读取状态、选择节点与资源、触发复核/上云", "不直接给出轴承故障类别"),
            ("边缘节点 E1", "振动预处理、时域/频域特征、轻量初检", "不强制处理所有数据"),
            ("边缘节点 E2", "独立模型或较保守阈值复核", "不在无触发条件时重复计算"),
            ("边缘节点 E3", "多传感融合、规则一致性检查或备用算力", "不替代云端复杂诊断"),
            ("云端", "冲突仲裁、复杂模型、历史趋势、模型更新与归档", "不作为所有数据的默认必经节点"),
        ],
        [1800, 4200, 3360],
        font_size=8.8,
    )
    add_callout(document, "路由原则", "发送器统一向 `task/raw` 发布；调度器订阅原始任务并根据策略向 `edge/{node_id}/task` 发布；边缘结果通过 `task/result` 回传；只有满足上云条件时才向 `cloud/diagnosis/task` 发布。")

    document.add_heading("4.1 节点异构性的表达", level=2)
    add_paragraph(document, "多个边缘节点不是同一设备的简单复制。节点可在 CPU 容量、当前负载、模型版本、特征能力和安全阈值上存在差异。调度状态中的 CR_t 与 R_t 描述瞬时资源和链路，节点能力标签则用于构建可行节点集合：不具备对应模型或最小资源的节点在决策前被屏蔽。这样可避免 Actor 选择物理上无法执行该子任务的节点。")

    document.add_heading("4.2 弱网自治边界", level=2)
    add_paragraph(document, "当边缘到云端通信波动或中断时，边缘节点继续完成基础安全判断，不等待云端才决定是否告警。对于疑似严重异常，可先执行本地保护动作或保守告警，同时缓存完整原始窗口；链路较弱但仍可用时，优先上传 RMS、Peak、峭度、频带能量、置信度和设备工况等特征摘要；链路恢复后再补传原始数据。该逻辑属于业务安全规则，不改变论文以完成时间为核心的奖励定义。")


def add_section_5(document):
    document.add_heading("5. 轴承数据采集、封装与通信流程", level=1)
    document.add_heading("5.1 时间与数据参数", level=2)
    add_table(
        document,
        ["参数", "设定值", "在调度中的含义"],
        [
            ("振动有效采样率", "16 kHz", "决定振动数据量和频域分析上限"),
            ("采集读取周期", "4 ms / 64 点", "连续写入发送缓冲区，不单独触发网络发送"),
            ("微批量发送周期", "50 ms / 800 点", "形成 packet 到达周期和网络负载"),
            ("边缘缓存窗口", "200 ms / 3200 点", "形成一次稳定的滑动诊断窗口"),
            ("边缘分析步长", "50 ms", "每个设备每秒最多形成 20 个新诊断窗口"),
            ("电流", "批次末最新瞬时值", "作为运行工况辅助量，不做电流波形诊断"),
            ("温度/转速/负载", "批次末最新有效值", "用于工况归一化和规则一致性判断"),
        ],
        [2400, 2100, 4860],
        font_size=9.2,
    )

    document.add_heading("5.2 建议的可扩展消息信封", level=2)
    add_paragraph(document, "现有 JSON 可以保留轴承字段，但应把“消息信封”和“数据载荷”分离。信封固定包含 packet_id、task_id、device_id、sensor_id、序号、起止时间、数据类型、编码、采样率、样本数、完整性校验和优先级；载荷通过 `payloads[]` 描述振动数组、标量、特征、图像或视频片段。新增传感器时只增加载荷描述，不改变调度器的基本解析框架。")
    add_bullet(document, "`packet_id` 用于消息去重；`batch_sequence_number` 用于检测丢包、乱序和重复。")
    add_bullet(document, "`task_id` 将同一设备、同一 200 ms 窗口中的多个 packet 关联起来。")
    add_bullet(document, "`data_type`、`encoding`、`shape` 和 `sample_rate` 让边缘节点判断是否具备处理能力。")
    add_bullet(document, "`quality_flags` 记录缺采样、时间漂移、校验失败和传感器离线，供调度器构建状态与可行节点。")

    document.add_heading("5.3 MQTT 主题和消息方向", level=2)
    add_number(document, "采集发送器向 `task/raw/{device_id}` 发布 50 ms 微批量，QoS 根据可接受重复和丢失程度配置。")
    add_number(document, "调度器订阅原始主题，完成校验、任务聚合、DAG 构造和 List_seq 生成。")
    add_number(document, "调度器将当前子任务发送到 `edge/{node_id}/task`，消息包含任务类型、输入引用、最晚开始时间和资源请求。")
    add_number(document, "边缘节点在 `edge/{node_id}/result` 返回特征、类别、置信度、耗时、实际 CPU 和质量标志。")
    add_number(document, "调度器依据结果决定结束、复核或向 `cloud/diagnosis/task` 升级，并在最终主题发布诊断结论。")


def add_section_6(document):
    document.add_heading("6. 轴承检测任务的 DAG 建模与排序", level=1)
    add_paragraph(document, "一次 200 ms 轴承诊断窗口被拆分为下列子任务。可选复核和云端任务采用条件门控：其前驱关系预先定义，但只有触发条件成立时才进入就绪队列；未触发时记录为跳过，不占用计算资源。")
    add_table(
        document,
        ["编号", "子任务", "主要输入", "直接前驱", "可执行位置"],
        [
            ("T1", "接收、去重与完整性校验", "4 个微批量及元数据", "无", "任一通用边缘节点"),
            ("T2", "振动去趋势、滤波与窗口化", "3200 点振动", "T1", "具备信号处理能力的边缘节点"),
            ("T3", "工况标量清洗与归一化", "电流、温度、转速、负载", "T1", "任一通用边缘节点"),
            ("T4", "时域/频域特征提取", "T2 输出", "T2", "振动分析边缘节点"),
            ("T5", "多传感融合与工况校正", "T3、T4 输出", "T3、T4", "融合型边缘节点"),
            ("T6", "轻量故障初检", "融合特征", "T5", "部署轴承模型的边缘节点"),
            ("T7", "置信度与指标一致性评估", "初检结果及质量标志", "T6", "调度器/通用边缘节点"),
            ("T8", "独立边缘复核（条件）", "窗口或关键特征", "T7", "不同模型/阈值的边缘节点"),
            ("T9", "云端深度诊断（条件）", "初检、复核、历史与原始数据", "T7 或 T8", "云端"),
            ("T10", "告警、保护建议与结果归档", "最终诊断结论", "T7/T8/T9 的最终分支", "边缘控制端与云端归档"),
        ],
        [650, 2500, 2200, 1500, 2510],
        font_size=8.2,
    )

    document.add_heading("6.1 List_seq 的生成", level=2)
    add_paragraph(document, "在一个确定窗口内，T1 是入口任务；T2 与 T3 在 T1 后可并行就绪；T5 必须等待 T3 和 T4；T8、T9 只有条件成立时才解锁。算法 3.1 对所有当前就绪任务计算 rank_i、urgency_i 和 W_i，并选择 W_i 最大者加入 List_seq。")
    add_paragraph(document, "例如，在振动数据量较大且 T4 位于关键路径时，T2/T4 的 rank 可能高于 T3；当某个任务距离全局容忍时延更近时，urgency 会提高其排序权重。最终序列可能为 `T1 -> T2 -> T4 -> T3 -> T5 -> T6 -> T7 -> T10`，也可能因任务参数和时间预算不同而变化。该示例只说明拓扑合法性，不代表固定最优序列。")

    document.add_heading("6.2 排序与节点分配必须分开", level=2)
    add_paragraph(document, "W_i 只回答“下一个调度哪个就绪任务”。节点剩余资源和链路速率属于执行状态，交由 PER-DDPG 决定。如果把节点负载直接混入 W_i，既会改变论文算法 3.1 的含义，也会让任务顺序和执行位置耦合，难以分析究竟是哪一部分带来性能变化。")


def add_section_7(document):
    document.add_heading("7. PER-DDPG 节点选择与资源分配", level=1)
    document.add_heading("7.1 每个调度步的状态", level=2)
    add_paragraph(document, "对 List_seq 中当前子任务 I_t，状态由三类量组成。")
    add_bullet(document, "任务特征 I_t：输入数据量、计算密度、最小 CPU、满足前驱依赖后的最早就绪时间。")
    add_bullet(document, "资源状态 CR_t：每个可行边缘节点和云端的剩余 CPU，必要时通过可行性掩码排除模型不匹配或容量不足的节点。")
    add_bullet(document, "通信状态 R_t：当前数据所在位置到各候选节点的有效传输速率，用于计算发送完整窗口、降采样数据或特征摘要的传输时间。")

    document.add_heading("7.2 动作和约束", level=2)
    add_paragraph(document, "Actor 输出候选节点概率 A_t 和资源比例 F_t。系统在可行节点中选择概率最大的节点，并将 F_t 解释为目标节点当前剩余 CPU 的请求比例。执行前必须检查任务最小 CPU、节点物理容量和整个执行区间内的可用资源；若暂时不足则等待，若物理容量永远不足则禁止该节点。")
    add_callout(document, "资源审计", "日志同时记录请求比例、请求 CPU、实际分配 CPU 和实际比例。Actor 原始动作仍用于经验学习，环境转换与时延计算使用满足硬约束后的实际资源。")

    document.add_heading("7.3 奖励、经验回放与训练闭环", level=2)
    add_paragraph(document, "每执行一个子任务，环境更新传输完成时间、排队等待时间、执行时间和当前系统总完成时刻，并给出负的当前最大完成时间作为奖励。经验 `(s_t, a_t, r_t, s_{t+1})` 写入 PER 缓冲区，TD 误差越大的样本越容易被再次采样。Critic 学习动作长期价值，Actor 根据 Critic 梯度调整策略，目标网络软更新以降低训练震荡。")
    add_paragraph(document, "训练早期加入较大的探索噪声，让策略尝试不同节点与资源比例；随着回合推进，噪声线性减小，使策略逐步稳定。完成一个 DAG 只构成一个回合样本，必须在不同负载、网络和任务参数下重复训练，才能学习对状态变化具有适应性的调度策略。")

    document.add_heading("7.4 检测准确性与奖励函数的关系", level=2)
    add_paragraph(document, "论文奖励聚焦总完成时间。轴承场景中的模型置信度、传感器完整性和严重异常不能被忽略，但为了保持论文算法思想，不把它们擅自改造成新的加权奖励项。更严谨的做法是：把模型能力、最小资源和数据完整性作为动作可行性条件，把低置信度、指标矛盾和严重异常作为复核/上云门控条件；诊断准确率作为系统验收指标单独报告。")


def add_section_8(document):
    document.add_heading("8. 工业轴承检测端到端流程", level=1)
    add_number(document, "设备运行与连续采样：振动传感器以 16 kHz 采样；采集任务每 4 ms 写入 64 个点，同时更新电流、温度、转速和负载。")
    add_number(document, "微批量封装：发送器每 50 ms 取出 800 个未发送振动点，附加最新标量、序号、时间范围、质量标志和校验信息。")
    add_number(document, "原始任务发布：发送器向 MQTT 原始主题发布，不知道也不选择目标边缘节点。")
    add_number(document, "窗口聚合：调度器或入口边缘节点按 task_id 将 4 个微批量组织为 200 ms、3200 点诊断窗口，处理丢包、重复和乱序。")
    add_number(document, "DAG 构造与排序：根据窗口数据类型和可选分支生成子任务图，算法 3.1 计算 rank、urgency、W 并输出 List_seq。")
    add_number(document, "动态调度：对 List_seq 当前任务读取 I_t、CR_t、R_t，PER-DDPG 输出节点选择和 CPU 比例，调度器经 MQTT 将任务路由到目标边缘节点或云端。")
    add_number(document, "边缘执行：节点完成预处理、特征提取、融合或轻量推理，返回结果、置信度、质量标志、耗时和实际资源占用。")
    add_number(document, "复核判断：若初检置信度充分、指标一致且无关键异常，直接形成边缘结论；若低置信度、边界状态、指标矛盾或严重异常，触发第二边缘节点独立复核。")
    add_number(document, "冲突升级：若两次结果一致，采用复核后的边缘结论；若结果冲突、数据质量不足或需复杂历史模型，向云端上传特征或原始窗口进行深度诊断。")
    add_number(document, "弱网自治：云端不可达时，边缘按安全规则执行本地告警/保护，优先上传特征摘要并缓存原始窗口，恢复通信后补传。")
    add_number(document, "结果闭环：最终诊断、告警级别、保护建议和追溯信息写入结果主题及数据库；执行转移写入 PER 经验池，支持后续训练。")
    add_callout(document, "流程主线", "采集数据 -> 形成诊断窗口 -> 构造 DAG -> 生成 List_seq -> 按步选择节点和资源 -> 边缘初检 -> 条件复核 -> 冲突上云 -> 告警与归档 -> 经验反馈。")


def add_section_9(document):
    document.add_heading("9. 当前问题、风险与解决方案", level=1)
    add_table(
        document,
        ["问题", "风险或原因", "解决方案"],
        [
            ("接口只适配固定传感器", "新增声学、图像或其他设备时需要改代码", "采用固定消息信封 + `payloads[]` 描述数据类型、编码、形状与采样率；调度器按能力标签匹配节点"),
            ("发送器直接选节点", "采集端需要维护节点 IP，难扩容且无法全局负载均衡", "发送器只发布 `task/raw`；调度器集中读取状态并控制 MQTT 路由"),
            ("一条数据广播给全部节点", "重复计算、网络开销大，结果冲突难管理", "默认单节点初检；仅在低置信度、边界、矛盾或严重异常时触发第二节点"),
            ("节点职责与能力不清", "Actor 可能选择未部署对应模型的节点", "建立节点能力注册表和可行性掩码，容量/模型不满足时不进入候选集合"),
            ("网络波动或云端中断", "传输时延增加，边缘等待云端导致业务停摆", "边缘保留基础安全规则；特征优先上传；原始数据本地缓存并支持断点补传"),
            ("丢包、重复和乱序", "200 ms 窗口不完整导致误判", "使用序号、起止时间、样本数和校验值；缺失窗口打质量标志并决定重传/降级"),
            ("条件复核与固定 DAG 不一致", "可选任务可能被错误地提前执行", "将复核和云端分析建模为门控子任务，条件成立才进入就绪队列，未触发则跳过"),
            ("把 List_seq 当成全局最优", "忽视节点选择和动态资源状态", "明确 List_seq 只保证依赖合法并提供启发式顺序；节点与资源仍由 PER-DDPG 学习"),
            ("认为第一轮即可完成学习", "把一次策略输出误认为最优策略", "保留探索、经验回放、Critic 估值和多回合训练，并使用独立评估回合验证"),
            ("只看请求资源比例", "实际 CPU 可能因最小资源和并发约束被修正", "同时记录请求比例、请求 CPU、实际 CPU 和实际比例，并用实际值计算时延"),
            ("奖励只优化时间", "可能误解为算法同时自动保证诊断准确率", "准确性通过模型准入、数据完整性、复核门控和独立验收指标保障，不虚构额外论文奖励"),
            ("缺少可复现评价", "单次曲线或单随机种子结论不稳定", "至少使用多个训练种子，对 FLE、RO、DQN、DDPG、PER-DDPG 进行相同场景比较，并报告均值和标准差"),
        ],
        [2200, 2900, 4260],
        font_size=8.1,
    )


def add_section_10(document):
    document.add_heading("10. 实施与验证指标", level=1)
    add_paragraph(document, "验证分为通信链路、调度正确性、算法训练和工业诊断四层。工业诊断指标用于证明系统可用，但不与论文奖励函数混为一谈。")
    add_table(
        document,
        ["验证层", "核心指标", "通过标准或观察方式"],
        [
            ("采集与通信", "采样点数、批次周期、丢包/重复/乱序率、重连时间", "确认 4 ms/64 点、50 ms/800 点、200 ms/3200 点一致；异常可追溯"),
            ("DAG 与排序", "前驱合法性、入口/出口、List_seq 拓扑顺序、门控任务状态", "任一后继不得早于前驱；未触发分支不占用资源"),
            ("资源调度", "传输/等待/执行时延、makespan、节点利用率、请求与实际 CPU", "无超配；最小 CPU 满足；日志可复核"),
            ("强化学习", "回合奖励、评估 makespan、收敛波动、多种子均值和标准差", "训练与评估分离；smoke 只验证链路，不冒充性能结论"),
            ("复核机制", "复核触发率、结果一致率、冲突上云率、云端上传比例", "常规任务不重复计算；高风险样本可升级"),
            ("弱网自治", "断网本地告警成功率、特征摘要时延、补传完整率", "云端不可达时边缘仍能完成基础安全判断"),
            ("诊断效果", "准确率、召回率、漏报率、误报率、置信度校准", "按轴承故障类别和工况独立评估，不由调度奖励替代"),
        ],
        [1700, 3300, 4360],
        font_size=8.7,
    )

    document.add_heading("10.1 推荐实施顺序", level=2)
    add_number(document, "先冻结 50 ms 微批量接口和 200 ms 任务窗口，完成序号、时间戳、校验与质量标志。")
    add_number(document, "搭建 MQTT Broker、调度器、两个异构边缘节点和云端模拟节点，打通主题与回执。")
    add_number(document, "实现固定 DAG、门控分支和算法 3.1，先用规则动作验证每个任务时间与依赖。")
    add_number(document, "接入 PER-DDPG，验证动作可行性、资源约束、奖励、经验优先级和探索噪声。")
    add_number(document, "加入弱网、节点忙、低置信度、复核冲突和严重异常场景，进行多种子训练与独立评估。")


def add_section_11(document):
    document.add_heading("11. 结论", level=1)
    add_paragraph(document, "两份文档可以被统一为一条清晰的工业轴承检测链路：第二份文档提供采集频率、批量周期、缓存窗口和基础接口；第一份文档补充网络影响、MQTT 路由、边缘自治、复核和云端升级机制；PER-DDPG 论文则为具有依赖关系的子任务提供排序、节点选择和连续资源分配方法。")
    add_paragraph(document, "严谨实现时必须保持职责分离：算法 3.1 决定就绪任务顺序，PER-DDPG 决定当前任务的节点和资源，边缘模型负责轴承诊断，业务规则决定是否复核与上云。排序序列不是最优证明，第一轮动作也不是学习完成；算法需要在变化的任务、资源和网络状态中反复交互，利用 PER 和 Critic 反馈逐步改进策略。")
    add_paragraph(document, "去除中间计算层后，方案仍保留论文核心思想，同时更贴合当前工程边界。常规窗口在边缘闭环，疑难和冲突样本按需复核或上云，弱网时保持本地安全能力，最终形成“采集—调度—执行—复核—升级—反馈”的完整工业流程。")


def build_report(output_path: Path) -> None:
    document = Document()
    configure_styles(document)
    configure_section(document.sections[0])
    add_running_header_footer(document.sections[0])
    document.core_properties.title = "工业轴承检测场景下的 PER-DDPG 任务调度与资源分配方案"
    document.core_properties.subject = "工业轴承检测、DAG 任务排序、PER-DDPG、边缘计算"
    document.core_properties.author = "项目技术方案"
    document.core_properties.keywords = "工业轴承, PER-DDPG, DAG, MQTT, 边缘计算"

    add_title_block(document)
    add_section_1(document)
    add_section_2(document)
    add_section_3(document)
    add_section_4(document)
    add_section_5(document)
    add_section_6(document)
    add_section_7(document)
    add_section_8(document)
    add_section_9(document)
    add_section_10(document)
    add_section_11(document)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    build_report(args.output)
    print(args.output.resolve())


if __name__ == "__main__":
    main()

