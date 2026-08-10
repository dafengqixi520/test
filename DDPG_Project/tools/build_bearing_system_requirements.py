from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from build_bearing_per_ddpg_report import (
    BLUE,
    DARK_BLUE,
    EAST_ASIA_FONT,
    INK,
    LIGHT,
    MUTED,
    add_bullet,
    add_callout,
    add_number,
    add_page_number,
    add_paragraph,
    add_table,
    configure_section,
    configure_styles,
    set_run_font,
)


OUTPUT_NAME = "工业轴承滚动检测场景_代码实现需求方案.docx"


def add_header_footer(section):
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("工业轴承滚动检测场景 | 代码实现需求方案")
    set_run_font(run, size=9, bold=True, color=MUTED)
    add_page_number(section.footer.paragraphs[0])


def add_title_block(document):
    kicker = document.add_paragraph()
    kicker.paragraph_format.space_before = Pt(8)
    kicker.paragraph_format.space_after = Pt(4)
    set_run_font(kicker.add_run("软件需求规格说明书"), size=10.5, bold=True, color=BLUE)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    title.paragraph_format.keep_with_next = True
    set_run_font(
        title.add_run("工业轴承滚动检测场景\nPER-DDPG代码实现需求方案"),
        size=23,
        bold=True,
        color=INK,
    )

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    set_run_font(
        subtitle.add_run("真实MQTT链路 + 仿真采集与诊断 + 可替换硬件适配器"),
        size=12.5,
        color=MUTED,
    )

    metadata = [
        ("依据文档", "《工业轴承检测场景_PER-DDPG技术方案》"),
        ("业务主体", "发送器、调度器、边缘节点、云端节点"),
        ("完整流程", "采集—调度—执行—复核—升级—反馈"),
        ("文档状态", "经用户逐部分确认的实施需求基线"),
        ("编制日期", "2026年7月15日"),
    ]
    for label, value in metadata:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        set_run_font(paragraph.add_run(f"{label}："), bold=True)
        set_run_font(paragraph.add_run(value))

    rule = document.add_paragraph()
    rule.paragraph_format.space_before = Pt(8)
    rule.paragraph_format.space_after = Pt(12)
    p_pr = rule._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_requirement(document, req_id, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_together = True
    set_run_font(paragraph.add_run(f"{req_id} "), bold=True, color=DARK_BLUE)
    set_run_font(paragraph.add_run(text))
    return paragraph


def add_code_block(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.2)
    paragraph.paragraph_format.right_indent = Inches(0.2)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.keep_together = True
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT)
    p_pr.append(shd)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), "Consolas")
    r_fonts.set(qn("w:hAnsi"), "Consolas")
    r_fonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)


def section_1(document):
    document.add_heading("1. 文档目标与范围", level=1)
    add_callout(
        document,
        "目标",
        "建设一套能够实际运行和验收的工业轴承滚动检测演示系统，以真实MQTT消息链路串联仿真发送器、PER-DDPG调度器、多个异构边缘节点和云端节点，完整实现采集、调度、执行、复核、升级和反馈。",
    )
    add_requirement(document, "GEN-01", "系统必须只有四类业务主体：发送器、调度器、边缘节点和云端节点。MQTT Broker属于通信基础设施，不得被描述为第五个业务主体。")
    add_requirement(document, "GEN-02", "第一版必须真实实现MQTT发布订阅、多进程节点、DAG排序、PER-DDPG决策、条件复核、云端升级、断线缓存、经验反馈和Web展示。")
    add_requirement(document, "GEN-03", "传感器、轴承诊断模型、通信速率和节点资源采用可重复仿真，并通过接口允许后续替换为真实硬件与模型。")
    add_requirement(document, "GEN-04", "运行时不得包含雾节点、对应动作槽位、状态字段、API或前端图例。")
    add_requirement(document, "GEN-05", "本阶段只定义和验证软件需求，不直接执行真实PLC停机、急停或保护动作。")

    document.add_heading("1.1 范围外事项", level=2)
    add_bullet(document, "真实采集卡和PLC厂商驱动开发。")
    add_bullet(document, "经过工业认证的轴承诊断模型训练及准确率承诺。")
    add_bullet(document, "生产设备的直接安全控制和责任闭环。")
    add_bullet(document, "跨工厂高可用集群、云平台弹性扩缩和大规模设备管理。")


def section_2(document):
    document.add_heading("2. 总体架构与职责边界", level=1)
    add_table(
        document,
        ["业务主体", "内部功能", "唯一责任", "禁止承担"],
        [
            ("发送器", "仿真采集、连续缓冲、微批量、MQTT发布、断网补传", "产生并可靠发送轴承原始数据", "节点选择、故障诊断"),
            ("调度器", "窗口聚合、DAG、List_seq、PER-DDPG、复核、升级、反馈、Web", "决定任务顺序、执行节点和资源比例", "直接生成轴承故障类别"),
            ("边缘节点", "预处理、特征、融合、初检、复核、本地安全规则", "执行调度任务并报告诊断结果", "全局任务排序、策略训练"),
            ("云端节点", "冲突仲裁、复杂诊断、历史趋势和最终解释", "处理明确升级的疑难任务", "订阅和处理全部原始数据"),
        ],
        [1500, 3600, 2700, 1560],
        font_size=8.8,
    )
    add_paragraph(document, "多个边缘节点运行同一程序，通过node_id、CPU容量、能力标签、模型版本和阈值配置体现差异。第一版默认启动edge_1、edge_2和edge_3，其中复核节点必须与初检节点不同。")

    document.add_heading("2.1 四个启动入口", level=2)
    add_code_block(
        document,
        "python -m sender\n"
        "python -m scheduler\n"
        "python -m edge --node-id edge_1\n"
        "python -m cloud",
    )

    document.add_heading("2.2 六阶段追踪矩阵", level=2)
    add_table(
        document,
        ["阶段", "责任主体", "核心输入", "核心产物", "主要验收"],
        [
            ("采集", "发送器", "仿真传感器", "50 ms微批量", "点数、序号、补传"),
            ("调度", "调度器", "200 ms窗口、节点状态", "DAG、List_seq、动作", "拓扑、掩码、资源"),
            ("执行", "边缘节点", "窗口或前驱输出", "特征、初检结果", "幂等、时延、置信度"),
            ("复核", "调度器+第二边缘", "初检结果和触发原因", "独立复核结果", "不同节点、条件触发"),
            ("升级", "调度器+云端", "冲突/低置信度任务", "FINAL或PROVISIONAL", "断网补传、唯一终态"),
            ("反馈", "调度器", "状态、动作、奖励、下一状态", "PER经验和模型更新", "TD误差、训练隔离"),
        ],
        [900, 1650, 2350, 2350, 2110],
        font_size=8.6,
    )


def section_3(document):
    document.add_heading("3. 发送器需求", level=1)
    requirements = [
        ("SND-01", "仿真振动传感器必须以16 kHz生成连续float32数据，每4 ms形成64个连续采样点。"),
        ("SND-02", "发送器必须每50 ms取出800个未发送振动点，并附加批次结束时最新的电流、温度、转速和负载。"),
        ("SND-03", "发送器只能向bearing/raw/{device_id}发布，不得保存目标边缘节点地址或执行节点选择。"),
        ("SND-04", "每个批次必须包含唯一message_id、递增sequence、开始/结束时间、采样率、样本数、质量标志和CRC32。"),
        ("SND-05", "MQTT不可用时，未确认批次必须写入本地SQLite待发队列；恢复后按设备和sequence顺序补传。"),
        ("SND-06", "发送器必须同时支持虚拟时钟和墙钟模式；虚拟时钟用于确定性测试，墙钟模式不承诺4 ms硬实时。"),
        ("SND-07", "采集实现必须符合SensorAdapter接口，使仿真器可以被真实采集卡或PLC适配器替换。"),
    ]
    for req_id, text in requirements:
        add_requirement(document, req_id, text)

    document.add_heading("3.1 SensorAdapter接口契约", level=2)
    add_code_block(
        document,
        "read_frame() -> SensorFrame\n"
        "SensorFrame = {timestamp_ns, vibration[64], current, temperature, speed, load, quality_flags}\n"
        "health() -> {connected, last_sample_at, error_code}",
    )


def section_4(document):
    document.add_heading("4. 调度器需求", level=1)
    requirements = [
        ("SCH-01", "订阅原始批次，按device_id、sequence和message_id进行校验、去重、丢包检测及乱序整理。"),
        ("SCH-02", "每台设备维护200 ms滚动窗口；首4个批次形成3200点窗口，此后每收到一个新批次形成一个步长为50 ms的新窗口。"),
        ("SCH-03", "每个窗口生成唯一task_id，并创建固定T1—T10任务图。"),
        ("SCH-04", "T1—T7为必选任务，T8和T9为条件门控任务，T10为唯一最终出口。"),
        ("SCH-05", "严格按算法3.1计算rank_i、urgency_i、W_i并生成拓扑合法List_seq。"),
        ("SCH-06", "必须按List_seq逐任务调用PER-DDPG；排序、节点选择和资源分配不得合并为一个未经说明的规则。"),
        ("SCH-07", "候选节点使用固定槽位[edge_1, edge_2, edge_3, cloud_0]，离线或能力不匹配节点必须应用动作掩码。"),
        ("SCH-08", "Actor输出节点概率A_t和资源比例F_t，环境记录原始动作和约束后的实际执行结果。"),
        ("SCH-09", "每次派发必须记录请求比例、请求CPU、实际CPU、有效比例、传输时间、等待时间和执行时间。"),
        ("SCH-10", "必须依据T7结果唯一地选择正常结束、第二边缘复核或云端升级。"),
        ("SCH-11", "训练模式允许更新Actor/Critic；推理模式只能使用冻结检查点并记录经验。"),
        ("SCH-12", "节点拒绝、任务超时或应用级确认缺失时最多重试3次，之后写入dead_letters并执行规定降级流程。"),
    ]
    for req_id, text in requirements:
        add_requirement(document, req_id, text)


def section_5(document):
    document.add_heading("5. 边缘节点需求", level=1)
    requirements = [
        ("EDG-01", "边缘节点必须通过配置声明node_id、逻辑CPU容量、能力标签、模型版本、阈值和本地缓存路径。"),
        ("EDG-02", "节点必须对task_id、subtask_id和attempt执行幂等检查，相同有效派发不得重复计算。"),
        ("EDG-03", "节点必须能够执行T1校验、T2预处理、T3归一化、T4特征、T5融合、T6初检及T7一致性判断。"),
        ("EDG-04", "第一版特征至少包含RMS、Peak、峰峰值、峭度和频带能量。"),
        ("EDG-05", "诊断输出必须包含类别、异常分数、置信度、质量标志、模型版本、开始/完成时间和实际资源。"),
        ("EDG-06", "执行T8的复核节点必须不同于初检节点，并使用不同模型版本或不同阈值配置。"),
        ("EDG-07", "云端不可用时，边缘节点必须执行本地安全规则、输出临时结论并缓存待升级数据。"),
        ("EDG-08", "第一版将F_t实现为逻辑资源预留、并发限制和时延模拟，不得宣称精确控制物理GHz。"),
        ("EDG-09", "节点必须每1秒发布心跳，包含可用CPU、队列长度、能力、模型版本、健康状态和最后活动时间。"),
    ]
    for req_id, text in requirements:
        add_requirement(document, req_id, text)

    document.add_heading("5.1 诊断模型接口", level=2)
    add_code_block(
        document,
        "predict(window, features, context) -> DiagnosisResult\n"
        "DiagnosisResult = {class_name, anomaly_score, confidence, feature_votes, model_version, explanation}",
    )


def section_6(document):
    document.add_heading("6. 云端节点需求", level=1)
    requirements = [
        ("CLD-01", "云端只能处理调度器明确升级的任务，不得订阅并默认处理全部原始数据。"),
        ("CLD-02", "升级输入必须包含初检、复核、特征摘要、质量标志、触发原因、历史引用及必要原始窗口。"),
        ("CLD-03", "云端输出必须包含最终类别、置信度、冲突解释、建议动作、模型版本和处理时延。"),
        ("CLD-04", "第一版云端诊断使用可控仿真模型，不得据此声明真实工业诊断准确率。"),
        ("CLD-05", "第一版保存最终诊断和历史趋势；云端模型训练及模型自动下发不属于强制范围。"),
    ]
    for req_id, text in requirements:
        add_requirement(document, req_id, text)


def section_7(document):
    document.add_heading("7. MQTT通信与消息契约", level=1)
    add_table(
        document,
        ["主题", "方向", "用途", "QoS"],
        [
            ("bearing/raw/{device_id}", "发送器 -> 调度器", "50 ms原始微批量", "1"),
            ("scheduler/edge/{node_id}/task", "调度器 -> 边缘", "子任务派发或复核", "1"),
            ("edge/{node_id}/ack", "边缘 -> 调度器", "应用级接收确认", "1"),
            ("edge/{node_id}/result", "边缘 -> 调度器", "子任务和诊断结果", "1"),
            ("scheduler/cloud/task", "调度器 -> 云端", "升级任务", "1"),
            ("cloud/ack", "云端 -> 调度器", "应用级接收确认", "1"),
            ("cloud/result", "云端 -> 调度器", "云端最终诊断", "1"),
            ("bearing/final/{device_id}", "调度器 -> 消费者", "唯一最终结果", "1"),
            ("node/{node_id}/heartbeat", "边缘/云端 -> 调度器", "资源和健康状态", "0"),
            ("system/dead-letter", "调度器 -> 审计", "无法处理的消息", "1"),
        ],
        [3500, 2100, 2800, 960],
        font_size=8.4,
    )
    add_requirement(document, "MQT-01", "任务和结果使用QoS 1，同时必须通过message_id和业务幂等键消除重复处理。")
    add_requirement(document, "MQT-02", "MQTT确认只代表Broker传输，边缘和云端必须额外返回应用级ack。")
    add_requirement(document, "MQT-03", "原始振动采用MessagePack信封和little-endian float32二进制载荷；控制、心跳和诊断结果使用UTF-8 JSON。")
    add_requirement(document, "MQT-04", "每条消息必须包含schema_version、message_id、trace_id、task_id、subtask_id、device_id、source_id、target_id、sequence、attempt、created_at_ns、deadline_at_ns、message_type、payload_encoding、quality_flags、model_version和crc32。")


def section_8(document):
    document.add_heading("8. T1—T10任务定义", level=1)
    add_table(
        document,
        ["任务", "名称", "输入", "输出", "依赖/门控"],
        [
            ("T1", "完整性校验", "4个连续微批量", "ValidatedWindow、质量标志", "入口任务"),
            ("T2", "振动预处理", "3200点振动", "滤波窗口", "T1"),
            ("T3", "工况归一化", "电流、温度、转速、负载", "NormalizedScalars", "T1，可与T2并行"),
            ("T4", "特征提取", "T2输出", "时域/频域特征", "T2"),
            ("T5", "多传感融合", "T3、T4输出", "FusedFeatures", "T3和T4全部完成"),
            ("T6", "轻量初检", "T5输出", "初检DiagnosisResult", "T5"),
            ("T7", "一致性判断", "初检和质量标志", "ACCEPT/REVIEW/CLOUD", "T6"),
            ("T8", "独立复核", "窗口、特征、初检结果", "复核DiagnosisResult", "条件门控，T7"),
            ("T9", "云端诊断", "初检、复核、历史、必要原始数据", "云端最终诊断", "条件门控，T7或T8"),
            ("T10", "告警与归档", "当前有效最终诊断", "FINAL或PROVISIONAL", "唯一出口"),
        ],
        [600, 1800, 2600, 2600, 1760],
        font_size=8.2,
    )
    add_requirement(document, "DAG-01", "任一后继任务只有在全部必需前驱完成后才能进入就绪队列。")
    add_requirement(document, "DAG-02", "未触发的T8或T9必须记录为SKIPPED，不得占用资源或生成虚假经验。")
    add_requirement(document, "DAG-03", "T10只能成功执行一次；同一task_id不得产生多个有效最终结果。")


def section_9(document):
    document.add_heading("9. PER-DDPG算法要求", level=1)
    add_callout(
        document,
        "固定结构",
        "第一版固定3个边缘槽位和1个云端槽位，节点顺序为[edge_1, edge_2, edge_3, cloud_0]。新增超过3个边缘节点会改变网络输入输出结构，必须重新训练。",
    )
    add_table(
        document,
        ["组成", "定义", "第一版维度"],
        [
            ("I_t", "数据量、计算密度、最小CPU、最早就绪时间", "4"),
            ("CR_t", "3个边缘节点和云端剩余逻辑CPU", "4"),
            ("R_t", "当前数据位置到4个候选节点的有效速率", "4"),
            ("状态s_t", "I_t + CR_t + R_t", "12"),
            ("A_t", "4个候选节点的选择概率", "4"),
            ("F_t", "目标节点剩余资源请求比例", "1"),
            ("动作a_t", "A_t + F_t", "5"),
        ],
        [1800, 5660, 1900],
        font_size=9,
    )
    add_requirement(document, "ALG-01", "通用状态维度为2K+6，动作维度为K+2；K表示模型支持的最大边缘槽位数。")
    add_requirement(document, "ALG-02", "离线节点槽位保留但CR和R设为0；能力不匹配节点通过掩码排除，并对剩余概率重新归一化。")
    add_requirement(document, "ALG-03", "requested_cpu=available_cpu×F_t；allocated_cpu=max(min_cpu, requested_cpu)，且不得超过执行区间可用容量。")
    add_requirement(document, "ALG-04", "每步奖励保持为当前最大完成时刻的负值，不增加准确率、告警等级、费用或能耗权重。")
    add_requirement(document, "ALG-05", "每个已执行子任务形成一条(s,a,r,s_next,done)经验；T10完成时done=True。")
    add_requirement(document, "ALG-06", "PER继续以TD误差生成优先级，LD-Noise在完整训练周期内线性衰减。")
    add_requirement(document, "ALG-07", "检查点必须记录schema_version、K、state_dim、action_dim和模型来源；结构不兼容时拒绝加载。")
    add_requirement(document, "ALG-08", "训练、验证和测试随机种子分离，正式算法比较至少运行5个独立训练种子。")


def section_10(document):
    document.add_heading("10. 复核、升级与弱网规则", level=1)
    add_requirement(document, "REV-01", "初检置信度低于配置阈值时触发第二边缘节点复核。")
    add_requirement(document, "REV-02", "异常分数位于正常/异常边界区间时触发复核。")
    add_requirement(document, "REV-03", "至少两项特征规则相互矛盾时触发复核。")
    add_requirement(document, "REV-04", "严重异常必须提高业务优先级并触发高优先级复核；没有可用复核节点时直接申请云端。")
    add_requirement(document, "REV-05", "初检与复核类别不同、复核后仍低置信度或数据质量不足时必须升级云端。")
    add_requirement(document, "REV-06", "云端不可用时必须形成PROVISIONAL临时结论，保存待升级数据；恢复后补传并更新为FINAL。")
    add_requirement(document, "REV-07", "置信度阈值、边界区间、严重异常阈值和矛盾票数必须配置化，不得散落硬编码。")

    document.add_heading("10.1 任务状态机", level=2)
    add_code_block(
        document,
        "COLLECTING -> READY -> SCHEDULED -> RUNNING\n"
        "RUNNING -> REVIEW_REQUIRED -> REVIEWING\n"
        "REVIEWING -> CLOUD_REQUIRED -> CLOUD_RUNNING\n"
        "RUNNING/REVIEWING/CLOUD_RUNNING -> FINALIZED\n"
        "任意执行状态 -> RETRYING -> FAILED\n"
        "云端不可达 -> PROVISIONAL -> PENDING_SYNC -> FINALIZED",
    )
    add_requirement(document, "STA-01", "每次状态变化必须保存时间、原因、来源组件、trace_id和版本号。")
    add_requirement(document, "STA-02", "非法状态跳转必须被拒绝并进入审计日志，不得静默修正。")


def section_11(document):
    document.add_heading("11. 数据存储与恢复", level=1)
    add_table(
        document,
        ["存储对象", "所有者", "用途"],
        [
            ("sender_spool", "发送器", "断网待发批次、确认状态和补传顺序"),
            ("messages", "调度器", "消息去重、序号、重试和应用级确认"),
            ("windows", "调度器", "200 ms窗口、质量标志和载荷引用"),
            ("tasks/subtasks", "调度器", "总任务与T1—T10状态、依赖和时延"),
            ("node_status", "调度器", "心跳、资源、能力和模型版本"),
            ("dispatches", "调度器", "Actor动作、掩码、请求及实际资源"),
            ("diagnoses", "调度器", "初检、复核、云端和最终结果"),
            ("transitions", "调度器", "PER经验、TD误差和优先级"),
            ("dead_letters", "调度器", "超过重试次数或格式无效的消息"),
            ("model_registry", "调度器/云端", "模型结构、版本、来源和校验值"),
        ],
        [2200, 1800, 5360],
        font_size=8.7,
    )
    add_requirement(document, "DAT-01", "第一版使用SQLite；组件重启后必须从持久化状态恢复未完成任务，不得只依赖内存。")
    add_requirement(document, "DAT-02", "原始窗口默认保留24小时，特征、诊断、审计和训练经验默认保留30天，均允许配置。")
    add_requirement(document, "DAT-03", "日志默认只记录原始数据引用、摘要和校验值，不输出完整敏感振动数组。")


def section_12(document):
    document.add_heading("12. 调度器Web控制台", level=1)
    items = [
        ("UI-01", "显示当前任务位于采集、调度、执行、复核、升级或反馈阶段。"),
        ("UI-02", "显示每台设备50 ms批次、200 ms窗口、丢包、乱序和补传状态。"),
        ("UI-03", "显示T1—T10 DAG、List_seq、前驱关系以及门控任务的EXECUTED或SKIPPED状态。"),
        ("UI-04", "显示Actor节点概率、动作掩码、目标节点、请求资源和实际资源。"),
        ("UI-05", "显示初检、复核、冲突原因、云端升级及FINAL/PROVISIONAL结论。"),
        ("UI-06", "显示传输、等待、执行时延、makespan、节点利用率和队列。"),
        ("UI-07", "显示PER缓冲区、TD误差、优先级、LD-Noise和训练曲线。"),
        ("UI-08", "显示Broker断线、节点离线、重试、死信和待补传告警。"),
    ]
    for req_id, text in items:
        add_requirement(document, req_id, text)


def section_13(document):
    document.add_heading("13. 非功能与安全要求", level=1)
    requirements = [
        ("NFR-01", "Broker地址、节点能力、阈值、模型路径、保留时间、随机种子和运行模式全部配置化。"),
        ("NFR-02", "每条消息、派发、诊断和状态变化必须共享trace_id，支持端到端追踪。"),
        ("NFR-03", "相同task_id只能形成一个有效最终结果；重复请求必须返回已有结果或幂等确认。"),
        ("NFR-04", "仿真必须支持固定随机种子，正常、低置信度、冲突、严重异常和断网场景可重复。"),
        ("NFR-05", "第一版默认用于本机或受控局域网；生产部署前必须增加TLS、账号、ACL、密钥轮换和审计。"),
        ("NFR-06", "模型、配置、数据库和实验结果必须记录生成时间、版本、来源和SHA-256。"),
        ("NFR-07", "单设备20批次/秒连续运行10分钟时，不得出现未解释消息丢失、重复终态或任务泄漏。"),
        ("NFR-08", "正常基线场景下，完整诊断任务应在配置的5秒全局容忍时延内完成；该值不代表工业认证指标。"),
    ]
    for req_id, text in requirements:
        add_requirement(document, req_id, text)


def section_14(document):
    document.add_heading("14. 测试与验收矩阵", level=1)
    add_table(
        document,
        ["编号", "场景", "输入/故障", "预期结果"],
        [
            ("E2E-01", "正常路径", "高置信度、数据完整", "初检后直接T10，仅一个FINAL"),
            ("E2E-02", "低置信度", "初检置信度低", "第二边缘复核，一致后结束"),
            ("E2E-03", "结果冲突", "初检与复核类别不同", "必须升级云端"),
            ("E2E-04", "严重异常", "严重异常分数", "高优先级复核；无节点则云端"),
            ("E2E-05", "边缘离线", "一个节点心跳超时", "动作掩码排除并改选节点"),
            ("E2E-06", "云端断线", "升级时云端不可达", "PROVISIONAL、缓存、恢复后FINAL"),
            ("E2E-07", "消息重复", "QoS 1重复投递", "不重复执行、不重复终态"),
            ("E2E-08", "乱序/丢失", "缺序号或乱序批次", "窗口质量异常，不静默正常诊断"),
            ("E2E-09", "Broker重启", "运行中重启Broker", "待发队列和任务可恢复"),
            ("E2E-10", "资源不足", "目标节点CPU不足", "等待或换节点，不超配"),
            ("E2E-11", "DAG正确性", "正常及门控分支", "前驱先于后继，未触发为SKIPPED"),
            ("E2E-12", "学习有效性", "缓冲区达到Batch Size", "Actor参数变化，PER优先级有限"),
            ("E2E-13", "结构检查", "配置、API、状态和前端", "不存在中间层运行节点"),
            ("E2E-14", "持续运行", "20批次/秒，10分钟", "无未解释丢失、重复终态和泄漏"),
        ],
        [900, 1700, 3000, 3760],
        font_size=8,
    )
    add_requirement(document, "TST-01", "单元测试覆盖采样、批量、窗口、CRC、DAG、排序、掩码、资源约束、状态机、去重和模型接口。")
    add_requirement(document, "TST-02", "集成测试必须使用真实Broker和独立进程，不得只调用进程内函数模拟MQTT。")
    add_requirement(document, "TST-03", "端到端测试必须从发送器开始，以bearing/final主题和持久化终态作为完成证据。")
    add_requirement(document, "TST-04", "Smoke测试只证明链路和输出结构，不得作为算法性能或诊断准确率结论。")


def section_15(document):
    document.add_heading("15. 建议代码结构", level=1)
    add_code_block(
        document,
        "bearing_per_ddpg_system/\n"
        "|-- common/       # 消息模型、MQTT、配置、日志和校验\n"
        "|-- sender/       # 采集、微批量和断网缓存\n"
        "|-- scheduler/    # 窗口、DAG、PER-DDPG、复核、反馈和Web\n"
        "|-- edge/         # 预处理、特征、初检和本地自治\n"
        "|-- cloud/        # 冲突仲裁、深度诊断和历史\n"
        "|-- algorithm/    # Actor、Critic、PER、LD-Noise和Runner\n"
        "|-- configs/      # Broker、节点、阈值、模型和模式\n"
        "|-- tests/        # unit、integration和e2e\n"
        "|-- scripts/      # 启动、训练、故障注入和验收\n"
        "`-- output/       # 模型、数据库、日志和实验结果",
    )
    add_callout(document, "边界说明", "common、algorithm、configs、tests和scripts是共享库或工程目录，不是新的业务服务；系统仍然只有四类业务主体。")


def section_16(document):
    document.add_heading("16. 实施阶段与阶段完成条件", level=1)
    stages = [
        ("阶段1：接口冻结", "冻结主题、字段、T1—T10、状态机和配置", "契约测试通过，字段无歧义"),
        ("阶段2：采集链路", "实现采样、微批量、MQTT和待发队列", "SND-01至SND-07及断线测试通过"),
        ("阶段3：基础调度执行", "窗口、DAG、List_seq和规则策略", "正常路径从发送器运行到FINAL"),
        ("阶段4：复核升级", "第二边缘复核、冲突上云和弱网补传", "E2E-02至E2E-06通过"),
        ("阶段5：PER-DDPG适配", "无中间层状态动作、掩码、资源和训练", "维度、约束、PER和模型版本测试通过"),
        ("阶段6：反馈与Web", "经验持久化、训练曲线和全流程界面", "六阶段、DAG、决策和异常可见"),
        ("阶段7：故障验收", "Broker、节点、网络、消息和资源故障注入", "14项E2E场景全部通过"),
    ]
    add_table(document, ["阶段", "主要工作", "阶段完成条件"], stages, [2100, 3660, 3600], font_size=8.8)


def section_17(document):
    document.add_heading("17. 难点、限制与不可承诺项", level=1)
    add_table(
        document,
        ["事项", "第一版处理方式", "不能承诺的内容"],
        [
            ("4 ms采集周期", "虚拟时钟保证点数和逻辑周期，墙钟模式尽力调度", "Windows/Python工业硬实时"),
            ("轴承诊断", "可控仿真模型和可替换ModelAdapter", "无真实数据时的工业准确率"),
            ("CPU资源比例", "逻辑预留、并发限制和时延模拟", "精确分配物理GHz"),
            ("MQTT可靠性", "QoS 1、持久队列、应用ack和幂等去重", "协议层严格Exactly Once"),
            ("网络速率", "按时间戳、字节数和故障注入估算有效速率", "底层物理带宽精确测量"),
            ("生产在线学习", "仿真训练，生产推理默认冻结模型", "未经验证的在线策略安全性"),
            ("工业保护动作", "只输出模拟告警和保护建议", "直接驱动急停或PLC"),
            ("动态扩容", "固定3个边缘槽位，槽位内节点可离线恢复", "超过槽位数仍复用旧检查点"),
        ],
        [2100, 3900, 3360],
        font_size=8.8,
    )


def section_18(document):
    document.add_heading("18. 最终完成定义", level=1)
    add_number(document, "四类业务主体能够以独立进程启动，并通过真实MQTT Broker通信。")
    add_number(document, "采集参数、窗口、T1—T10、List_seq和PER-DDPG状态动作全部符合本需求。")
    add_number(document, "正常、复核、升级、弱网和反馈路径均可从Web与持久化记录追踪。")
    add_number(document, "14项端到端场景、单元测试和真实Broker集成测试全部通过。")
    add_number(document, "系统不存在重复最终结果、未解释消息丢失、资源超配或非法状态跳转。")
    add_number(document, "模型、配置、数据库和实验输出具有版本、来源、时间和校验值。")
    add_number(document, "文档列出的困难项得到如实标注，不将仿真能力包装为工业认证能力。")
    add_callout(document, "交付边界", "本需求文档通过不等于系统代码已经实现。系统开发应按阶段1至阶段7逐步进行，每个阶段均以对应测试证据作为完成条件。")


def build_document(output_path: Path):
    document = Document()
    configure_styles(document)
    configure_section(document.sections[0])
    add_header_footer(document.sections[0])
    document.core_properties.title = "工业轴承滚动检测场景PER-DDPG代码实现需求方案"
    document.core_properties.subject = "采集、调度、执行、复核、升级、反馈完整工业流程的软件需求"
    document.core_properties.author = "项目需求基线"
    document.core_properties.keywords = "工业轴承, PER-DDPG, MQTT, DAG, 边缘计算, 软件需求"

    add_title_block(document)
    section_1(document)
    section_2(document)
    section_3(document)
    section_4(document)
    section_5(document)
    section_6(document)
    section_7(document)
    section_8(document)
    section_9(document)
    section_10(document)
    section_11(document)
    section_12(document)
    section_13(document)
    section_14(document)
    section_15(document)
    section_16(document)
    section_17(document)
    section_18(document)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    build_document(args.output)
    print(args.output.resolve())


if __name__ == "__main__":
    main()
