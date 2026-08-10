from pathlib import Path
import math
import os

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "frontend_docs"
OUT.mkdir(parents=True, exist_ok=True)
IMAGE_PATH = OUT / "industrial_iot_frontend_data_flow.png"
DOCX_PATH = OUT / "industrial_iot_per_ddpg_frontend_guide.docx"

FONT_CANDIDATES = [
    r"C:\Windows\Fonts\msyh.ttc",
    r"C:\Windows\Fonts\simhei.ttf",
    r"C:\Windows\Fonts\simsun.ttc",
]
FONT_PATH = next((p for p in FONT_CANDIDATES if os.path.exists(p)), None)


def font(size):
    if FONT_PATH:
        return ImageFont.truetype(FONT_PATH, size)
    return ImageFont.load_default()


def draw_wrapped(draw, text, box, fnt, fill="#172033", gap=6):
    x1, y1, x2, _ = box
    max_width = x2 - x1
    lines = []
    for para in text.split("\n"):
        current = ""
        for ch in para:
            test = current + ch
            if draw.textbbox((0, 0), test, font=fnt)[2] <= max_width or not current:
                current = test
            else:
                lines.append(current)
                current = ch
        lines.append(current)
    y = y1
    for line in lines:
        draw.text((x1, y), line, font=fnt, fill=fill)
        bbox = draw.textbbox((0, 0), line, font=fnt)
        y += (bbox[3] - bbox[1]) + gap


def build_image():
    width, height = 1800, 1320
    img = Image.new("RGB", (width, height), "#f7fafc")
    draw = ImageDraw.Draw(img)
    colors = {
        "muted": "#5b677a",
        "blue": "#2563eb",
        "cyan": "#0891b2",
        "green": "#16a34a",
        "orange": "#ea580c",
        "red": "#dc2626",
        "purple": "#7c3aed",
        "line": "#94a3b8",
        "white": "#ffffff",
        "dark": "#0f172a",
    }

    def rounded_rect(xy, radius, fill, outline="#cbd5e1", line_width=2):
        draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=line_width)

    draw.text((60, 38), "工业物联网 PER-DDPG 前端数据与任务流程对应图", font=font(38), fill=colors["dark"])
    draw.text(
        (62, 92),
        "说明每张前端图表/表格的数据是在算法流程的哪个阶段产生、按什么粒度记录。",
        font=font(22),
        fill=colors["muted"],
    )

    stages = [
        (
            "1. 初始化/配置",
            "读取节点数量、CPU能力、训练模式、速度等。\n前端对应：顶部按钮、模式、速度、状态徽标。",
            colors["blue"],
        ),
        (
            "2. 生成 DAG 与子任务",
            "每个 Episode 开始时生成工业子任务、数据量、计算密度、依赖边。\n前端对应：DAG依赖图、子任务表基础字段。",
            colors["cyan"],
        ),
        (
            "3. 算法3.1 任务排序",
            "计算 rank、urgency、W_i，得到 List_seq。\n前端对应：DAG右侧调度序列、子任务表优先级列。",
            colors["purple"],
        ),
        (
            "4. PER-DDPG 单步决策",
            "策略网络根据状态选择卸载节点和CPU比例；同时产生探索噪声η。\n前端对应：探索噪声η衰减、训练日志。",
            colors["orange"],
        ),
        (
            "5. 环境执行与节点资源更新",
            "执行当前子任务，计算传输、等待、执行、完成时间，更新云/雾/边资源。\n前端对应：子任务表、计算节点卡片、DAG节点状态。",
            colors["green"],
        ),
        (
            "6. 奖励与训练更新",
            "每步获得奖励；经验池满批量后更新价值网络和策略网络；一轮结束记录累计奖励和完工时延。\n前端对应：奖励曲线、Makespan曲线、Loss。",
            colors["red"],
        ),
    ]

    x0, y0, box_w, box_h, gap = 60, 165, 520, 160, 48
    centers = []
    for idx, (title, desc, color) in enumerate(stages):
        row, col = divmod(idx, 2)
        x = x0 + col * 640
        y = y0 + row * (box_h + gap)
        rounded_rect((x, y, x + box_w, y + box_h), 18, colors["white"], color, 4)
        draw.ellipse((x + 20, y + 20, x + 58, y + 58), fill=color)
        draw.text((x + 31, y + 22), str(idx + 1), font=font(22), fill="white")
        draw.text((x + 74, y + 18), title, font=font(26), fill=colors["dark"])
        draw_wrapped(draw, desc, (x + 74, y + 58, x + box_w - 22, y + box_h - 16), font(16), colors["muted"], 5)
        centers.append((x + box_w / 2, y + box_h / 2))

    for a, b in [(0, 1), (1, 2), (2, 3), (3, 4), (4, 5)]:
        ax, ay = centers[a]
        bx, by = centers[b]
        if a % 2 == 0 and b == a + 1:
            start = (ax + box_w / 2 - 18, ay)
            end = (bx - box_w / 2 + 18, by)
        else:
            start = (ax, ay + box_h / 2 - 4)
            end = (bx, by - box_h / 2 + 4)
        draw.line((start, end), fill=colors["line"], width=4)
        ex, ey = end
        sx, sy = start
        angle = math.atan2(ey - sy, ex - sx)
        arrow_points = [
            (ex + 18 * math.cos(angle + math.pi * 0.82), ey + 18 * math.sin(angle + math.pi * 0.82)),
            (ex + 18 * math.cos(angle - math.pi * 0.82), ey + 18 * math.sin(angle - math.pi * 0.82)),
        ]
        draw.polygon([end] + arrow_points, fill=colors["line"])

    panel_x, panel_y, panel_w, panel_h = 60, 835, 1680, 405
    rounded_rect((panel_x, panel_y, panel_x + panel_w, panel_y + panel_h), 20, "#ffffff", "#cbd5e1", 2)
    draw.text((panel_x + 28, panel_y + 24), "前端图表/表格的数据粒度与来源", font=font(30), fill=colors["dark"])
    items = [
        ("奖励曲线", "横轴为轮次编号；纵轴为一轮完整DAG调度的累计奖励。来源：阶段6。"),
        ("Makespan曲线", "横轴为轮次编号；纵轴为该轮全部子任务最大完成时间，单位秒。来源：阶段6。"),
        ("探索噪声η", "横轴为调度步；纵轴为探索系数η。来源：阶段4，每个子任务决策前产生。"),
        ("Actor/Critic损失", "横轴为训练步；纵轴为价值网络和策略网络损失。来源：阶段6。"),
        ("DAG/List_seq", "节点是子任务，边是依赖和中间数据量；序列来自算法3.1排序。来源：阶段2-3。"),
        ("子任务/节点", "记录任务分配、CPU、传输/等待/执行/完成时间，以及节点CPU占用和队列。来源：阶段5。"),
    ]
    for idx, (name, desc) in enumerate(items):
        col, row = idx % 2, idx // 2
        x = panel_x + 30 + col * 820
        y = panel_y + 88 + row * 92
        draw.rounded_rectangle((x, y, x + 780, y + 72), radius=12, fill="#f8fafc", outline="#e2e8f0", width=1)
        draw.text((x + 18, y + 12), name, font=font(22), fill=colors["blue"])
        draw_wrapped(draw, desc, (x + 205, y + 12, x + 760, y + 62), font(15), colors["muted"], 3)

    img.save(IMAGE_PATH)


def set_cell(cell, text, bold=False, fill=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    run.font.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        tc_pr.append(shading)


def add_table(doc, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, value in enumerate(rows[0]):
        set_cell(table.rows[0].cells[idx], value, True, "E8EEF5")
    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell(cells[idx], value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths[: len(row.cells)]):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph("")


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color in [("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5"), ("Heading 3", 12, "1F4D78")]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("工业物联网 PER-DDPG 项目前端说明文档")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")
    subtitle = doc.add_paragraph(
        "面向云-雾-边协同任务卸载演示系统，说明主界面控件、各页面数据含义、图表横纵坐标，以及这些数据在算法任务流程中的产生阶段。"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead = doc.add_paragraph()
    lead.add_run(
        "核心理解：前端不是展示真实传感器训练模型，而是展示论文算法在工业物联网任务调度仿真中的状态、决策和训练指标。上方两张曲线按 Episode 统计，下面两张曲线按 Step 统计。"
    ).bold = True

    sections = [
        (
            "1. 主界面整体结构",
            "系统首页围绕一次 PER-DDPG 云-雾-边任务卸载训练展开。顶部是控制栏和状态信息，中间通过标签页切换不同观察视角：训练仪表盘、DAG依赖图、子任务表、计算节点、训练日志和实验对比。",
            [
                ["区域", "记录/展示内容", "对应算法阶段"],
                ["顶部控制栏", "训练、暂停、继续、单步、重置、模式、速度、状态徽标。", "运行控制阶段"],
                ["训练仪表盘", "奖励、Makespan、探索噪声、Loss和历史Episode表。", "训练统计与收敛观察"],
                ["DAG依赖图", "子任务依赖边、入口/出口任务、List_seq调度序列、rank/urgency/W_i。", "DAG生成与算法3.1排序"],
                ["子任务表", "每个子任务的数据量、计算密度、状态、分配节点、CPU、传输/等待/执行时间。", "环境执行与任务完成记录"],
                ["计算节点", "云/雾/边节点CPU总量、可用CPU、使用率、队列、已处理任务数。", "资源调度与节点状态更新"],
                ["实验对比", "本地复现实验中不同算法的平均Makespan、标准差、样本数、训练更新次数。", "算法效果对比"],
            ],
            [1.2, 3.5, 1.6],
        ),
        (
            "2. 顶部按钮与参数含义",
            "",
            [
                ["按钮/参数", "含义", "使用说明"],
                ["训练", "启动一次连续仿真/训练循环。", "在 Train 模式下，Actor动作会叠加探索噪声，并在经验池达到批量大小后更新网络。"],
                ["暂停", "暂停后台训练循环。", "暂停后页面保留当前DAG、子任务和节点状态。"],
                ["继续", "从暂停状态恢复训练。", "继续沿用当前runner和经验池。"],
                ["单步", "只执行一个子任务调度Step。", "适合观察每一步任务如何被分配到云/雾/边节点。"],
                ["重置", "重置仿真环境。", "会重新生成当前DAG和任务状态，历史训练曲线可能仍读取保存结果。"],
                ["模式 Train/Eval", "Train用于训练；Eval用于评估。", "Eval不加探索噪声，也不进行网络更新。"],
                ["速度", "控制每个Step之间的前端展示延迟。", "只影响展示节奏，不改变算法公式。"],
            ],
            [1.25, 1.75, 3.2],
        ),
        (
            "3. 训练仪表盘",
            "训练仪表盘用于观察策略学习过程和调度性能变化。它包含四张曲线和一个最近Episode表。",
            [
                ["图表/表格", "横坐标", "纵坐标/列", "数据产生阶段", "判断方向"],
                ["Episode奖励曲线", "Episode编号", "该Episode累计奖励。由于reward与负Makespan相关，数值越接近0越好。", "一轮完整DAG调度结束后汇总。", "整体上升更好。"],
                ["Makespan收敛曲线", "Episode编号", "Makespan(s)，即该轮所有子任务最大完成时间 max(FT_i)。", "Episode结束时统计。", "整体下降并稳定更好。"],
                ["探索噪声η衰减", "Step编号", "LD-Noise探索系数η。", "每个子任务决策前产生。", "线性下降是正常现象。"],
                ["Critic / Actor Loss", "Step编号", "Critic Loss为Q网络误差；Actor Loss为策略网络优化损失。", "经验池达到batch_size后，每次训练更新产生。", "不应爆炸或NaN；RL中不要求单调下降。"],
                ["最近Episode表", "表格行=Episode", "累计奖励、Makespan(s)。", "Episode结束后写入历史。", "用于查看最近若干轮稳定性。"],
            ],
            [1.25, 1.0, 1.9, 1.65, 1.2],
        ),
        (
            "4. DAG依赖图",
            "DAG依赖图展示一轮工业任务中的子任务结构。节点表示子任务 I_i，边表示前驱依赖和中间数据传输，右侧 List_seq 是算法3.1根据 rank、urgency 和 W_i 得到的调度顺序。",
            [
                ["字段", "含义"],
                ["节点 I0/I1/...", "工业子任务，例如振动信号窗口化、热图像预处理、声学异常特征等。"],
                ["边 from -> to", "任务依赖关系，后继任务必须等待前驱完成及中间数据传输。"],
                ["边标签 KB", "前驱到后继的中间数据量。"],
                ["rank", "静态优先级，反映任务及其后继链路的平均处理代价。"],
                ["urgency", "动态紧迫度，与当前虚拟时间和deadline相关。"],
                ["W_i", "综合调度优先级，由rank和urgency加权得到。"],
                ["List_seq", "最终任务执行顺序；PER-DDPG按该顺序逐个决定卸载节点和CPU比例。"],
            ],
            [1.35, 5.0],
        ),
        (
            "5. 子任务表",
            "子任务表按任务粒度记录每个子任务的输入属性、依赖关系和执行结果。它回答“这个任务被调度到哪里、花了多少时间、是否已经完成”。",
            [
                ["列名", "含义"],
                ["ID/名称", "子任务编号和工业场景语义名称。"],
                ["数据量(KB)", "任务输入数据大小。"],
                ["计算密度", "单位数据量所需CPU周期，影响执行时间。"],
                ["状态", "pending、ready、executing、completed等。"],
                ["分配节点", "Actor最终选择的edge/fog/cloud节点。"],
                ["分配CPU", "根据动作中的资源比例得到的CPU资源。"],
                ["开始->完成", "start_time到finish_time。"],
                ["传输/等待/执行", "输入传输时间 + 依赖传输时间、排队等待时间、实际计算执行时间。"],
                ["优先级W_i", "算法3.1计算的综合调度优先级。"],
                ["前驱->后继", "DAG依赖关系。"],
            ],
            [1.4, 5.0],
        ),
        (
            "6. 计算节点",
            "计算节点页面按资源节点展示云、雾、边三层资源状态。它回答“任务分配后资源是否拥塞、CPU是否被占用、队列是否增加”。",
            [
                ["字段", "含义"],
                ["节点ID", "cloud_0、fog_1...fog_M、edge_0...edge_N。"],
                ["节点类型", "cloud、fog、edge。"],
                ["CPU总量", "节点最大计算能力，单位GHz。"],
                ["可用CPU", "当前全局时间下可调度的剩余CPU。"],
                ["使用率", "CPU占用比例，用于观察资源压力。"],
                ["队列", "该节点已预约/排队的任务数量。"],
                ["已处理", "该节点累计处理过的任务数量。"],
            ],
            [1.4, 5.0],
        ),
        (
            "7. 实验对比",
            "实验对比页面读取本地实验输出文件，用于比较不同调度算法或配置的整体效果。该页面不代表单个Step的即时状态，而是多次实验样本的统计摘要。",
            [
                ["列名", "含义"],
                ["算法", "参与对比的调度/卸载策略名称。"],
                ["平均Makespan(s)", "该算法在样本集合上的平均完工总时延。越小越好。"],
                ["标准差", "Makespan波动程度。越小表示稳定性越好。"],
                ["样本数", "参与统计的Episode或实验次数。"],
                ["训练更新", "该算法训练过程中发生的网络更新次数；非学习算法可能为空。"],
            ],
            [1.4, 5.0],
        ),
    ]

    for heading, paragraph, rows, widths in sections:
        doc.add_heading(heading, level=1)
        if paragraph:
            doc.add_paragraph(paragraph)
        add_table(doc, rows, widths)

    doc.add_heading("8. 数据与任务流程阶段对应图", level=1)
    doc.add_paragraph(
        "下图把前端产生的数据映射到算法流程阶段。例如，奖励曲线中的每个点表示一个Episode编号对应的累计奖励，它是在“一轮完整DAG调度完成”后由环境奖励汇总得到的。"
    )
    doc.add_picture(str(IMAGE_PATH), width=Inches(6.5))
    caption = doc.add_paragraph("图1  前端数据与PER-DDPG工业物联网任务调度流程的对应关系")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("9. 使用时的读图建议", level=1)
    for item in [
        "奖励曲线和Makespan曲线按Episode观察趋势，不要用单个点判断是否收敛。",
        "探索噪声η按Step线性下降，平滑下降是正常的，不应理解为性能曲线。",
        "Critic/Actor Loss只有在经验池达到batch_size后才开始有真实训练意义；强化学习Loss通常会波动。",
        "DAG、子任务表、计算节点是同一轮调度的不同视角，应结合List_seq、分配节点和完成时间一起解释。",
    ]:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.2)
        paragraph.add_run("- ").bold = True
        paragraph.add_run(item)

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.text = "工业物联网 PER-DDPG 前端说明文档"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].font.size = Pt(9)
        footer.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_image()
    build_docx()
    print(IMAGE_PATH)
    print(DOCX_PATH)
