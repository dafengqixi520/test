from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document


REQUIRED_TERMS = [
    "发送器",
    "调度器",
    "边缘节点",
    "云端节点",
    "采集",
    "调度",
    "执行",
    "复核",
    "升级",
    "反馈",
    "16 kHz",
    "4 ms",
    "64",
    "50 ms",
    "800",
    "200 ms",
    "3200",
    "List_seq",
    "PER-DDPG",
    "MQTT",
    "状态s_t",
    "12",
    "动作a_t",
    "5",
    "COLLECTING",
    "FINALIZED",
    "PROVISIONAL",
    "PENDING_SYNC",
    "bearing/raw/{device_id}",
    "scheduler/edge/{node_id}/task",
    "scheduler/cloud/task",
    "bearing/final/{device_id}",
]

REQUIREMENT_PREFIX_COUNTS = {
    "GEN": 5,
    "SND": 7,
    "SCH": 12,
    "EDG": 9,
    "CLD": 5,
    "MQT": 4,
    "DAG": 3,
    "ALG": 8,
    "REV": 7,
    "STA": 2,
    "DAT": 3,
    "UI": 8,
    "NFR": 8,
    "E2E": 14,
    "TST": 4,
}


def collect_text(document):
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("path", type=Path)
    args = parser.parse_args()

    document = Document(args.path)
    text = collect_text(document)
    problems = []

    missing = [term for term in REQUIRED_TERMS if term not in text]
    if missing:
        problems.append(f"missing terms: {missing}")

    for prefix, expected_count in REQUIREMENT_PREFIX_COUNTS.items():
        found = sorted(set(re.findall(rf"\b{prefix}-\d{{2}}\b", text)))
        if len(found) != expected_count:
            problems.append(
                f"{prefix} expected {expected_count} unique IDs, found {len(found)}: {found}"
            )

    tasks = sorted(set(re.findall(r"\bT(?:10|[1-9])\b", text)))
    expected_tasks = [f"T{i}" for i in range(1, 11)]
    if set(tasks) != set(expected_tasks):
        problems.append(f"T1-T10 coverage mismatch: {tasks}")

    heading1 = [p.text for p in document.paragraphs if p.style.name == "Heading 1"]
    heading2 = [p.text for p in document.paragraphs if p.style.name == "Heading 2"]
    if len(heading1) != 18:
        problems.append(f"expected 18 Heading 1 sections, found {len(heading1)}")
    if len(document.tables) != 9:
        problems.append(f"expected 9 real tables, found {len(document.tables)}")

    header_text = " ".join(p.text for p in document.sections[0].header.paragraphs).strip()
    if "代码实现需求方案" not in header_text:
        problems.append("running header is missing or incorrect")
    if "PAGE" not in document.sections[0].footer._element.xml:
        problems.append("footer PAGE field is missing")

    business_subjects = ["发送器", "调度器", "边缘节点", "云端节点"]
    for subject in business_subjects:
        if text.count(subject) < 2:
            problems.append(f"business subject insufficiently specified: {subject}")

    print(f"paragraphs={len(document.paragraphs)}")
    print(f"heading1={len(heading1)} heading2={len(heading2)}")
    print(f"tables={len(document.tables)}")
    print(f"requirements={sum(REQUIREMENT_PREFIX_COUNTS.values())}")
    print(f"tasks={','.join(expected_tasks)}")

    if problems:
        for problem in problems:
            print(f"ERROR: {problem}")
        raise SystemExit(1)
    print("AUDIT PASS")


if __name__ == "__main__":
    main()

