from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


def extract(path: Path) -> None:
    document = Document(path)
    print(f"FILE: {path}")
    print("PARAGRAPHS")
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if text:
            print(f"P{index:03d} [{paragraph.style.name}] {text}")

    for table_index, table in enumerate(document.tables):
        print(f"TABLE {table_index} {len(table.rows)}x{len(table.columns)}")
        for row_index, row in enumerate(table.rows):
            cells = [cell.text.replace("\n", " / ").strip() for cell in row.cells]
            print(f"R{row_index}: " + " || ".join(cells))

    print(f"SECTIONS: {len(document.sections)}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("path", type=Path)
    args = parser.parse_args()
    extract(args.path)


if __name__ == "__main__":
    main()
